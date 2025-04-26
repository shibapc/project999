from docx import Document
from docx.shared import Pt, Mm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.enum.table import WD_ALIGN_VERTICAL
from datetime import datetime
import os
import json
import logging
from docx.image.exceptions import UnrecognizedImageError
from docx.oxml import OxmlElement

logger = logging.getLogger(__name__)

def load_materials_db() -> dict:
    """Загружает materials.json."""
    try:
        with open("materials.json", "r", encoding="utf-8") as f:
            return json.load(f)
    except FileNotFoundError:
        logger.error("Файл materials.json не найден")
        raise
    except json.JSONDecodeError:
        logger.error("Ошибка декодирования materials.json")
        raise

def set_cell_margins(cell, left=100, right=100, top=50, bottom=50):
    """Устанавливает отступы внутри ячейки таблицы."""
    tcPr = cell._element.tcPr
    if tcPr is None:
        tcPr = OxmlElement('w:tcPr')
        cell._element.append(tcPr)

    tcMar = tcPr.find(qn('w:tcMar'))
    if tcMar is None:
        tcMar = OxmlElement('w:tcMar')
        tcPr.append(tcMar)

    for margin, value in [('left', left), ('right', right), ('top', top), ('bottom', bottom)]:
        margin_elem = OxmlElement(f'w:{margin}')
        margin_elem.set(qn('w:w'), str(value))
        margin_elem.set(qn('w:type'), 'dxa')
        tcMar.append(margin_elem)

def set_cell_borders(cell):
    """Устанавливает границы ячейки таблицы."""
    tcPr = cell._element.tcPr
    if tcPr is None:
        tcPr = OxmlElement('w:tcPr')
        cell._element.append(tcPr)

    tcBorders = tcPr.find(qn('w:tcBorders'))
    if tcBorders is None:
        tcBorders = OxmlElement('w:tcBorders')
        tcPr.append(tcBorders)

    for border in ['top', 'left', 'bottom', 'right']:
        border_elem = tcBorders.find(qn(f'w:{border}'))
        if border_elem is None:
            border_elem = OxmlElement(f'w:{border}')
            tcBorders.append(border_elem)
        border_elem.set(qn('w:val'), 'single')
        border_elem.set(qn('w:sz'), '4')  # Толщина границы (4 = 0.5 pt)
        border_elem.set(qn('w:space'), '0')
        border_elem.set(qn('w:color'), '000000')  # Черный цвет

def create_commercial_proposal_docx(chat_id: int, user_data: dict, output_dir: str = "commercial_offer") -> str:
    """Создаёт коммерческое предложение в формате DOCX."""
    try:
        # Создаём директорию
        os.makedirs(output_dir, exist_ok=True)
        filename = os.path.join(output_dir, f"КП_{chat_id}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx")
        doc = Document()

        # Установка полей страницы (15 мм + 30 мм сверху для отступа 3 см)
        section = doc.sections[0]
        section.left_margin = Mm(15)
        section.right_margin = Mm(15)
        section.top_margin = Mm(15 + 30)  # Добавляем 30 мм к верхнему полю (итого 45 мм)
        section.bottom_margin = Mm(15)

        # Функция для установки шрифта Times New Roman
        def set_font(run, size, bold=False):
            run.font.name = 'Times New Roman'
            run.font.size = Pt(size)
            run.font.bold = bold
            run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')

        # Шапка: таблица для логотипа и реквизитов
        header_table = doc.add_table(rows=1, cols=2)
        header_table.autofit = False
        header_table.columns[0].width = Mm(93)  # Ширина для логотипа
        header_table.columns[1].width = Mm(100)  # Ширина для реквизитов

        # Логотип слева
        logo_cell = header_table.cell(0, 0)
        logo_paragraph = logo_cell.paragraphs[0]
        logo_run = logo_paragraph.add_run()
        logo_path = "utils/commercial_offer/soup.jpg"
        try:
            if os.path.exists(logo_path):
                logo_run.add_picture(logo_path, width=Mm(93)) 
            else:
                logger.warning(f"Файл логотипа не найден по пути: {logo_path}")
                logo_run.add_text("Логотип отсутствует")
                set_font(logo_run, 12)
        except UnrecognizedImageError as e:
            logger.error(f"Ошибка при загрузке логотипа: {str(e)}")
            logo_run.add_text("Логотип не поддерживается")
            set_font(logo_run, 12)
        except Exception as e:
            logger.error(f"Неизвестная ошибка при загрузке логотипа: {str(e)}")
            logo_run.add_text("Ошибка загрузки логотипа")
            set_font(logo_run, 12)

        # Реквизиты справа
        details_cell = header_table.cell(0, 1)
        details_paragraph = details_cell.paragraphs[0]
        details_paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        details_text = (
            "ООО «Студия Уникальных Проектов»\n"
            "ИНН/КПП 62710330360/669801001\n"
            "Юр. адрес ул. 620131 г. Екатеринбург\n"
            "ул. Фурманова, дом 19, корпус 1 помещ. 2\n"
            "Факт. адрес ул. 620131 г. Екатеринбург\n"
            "ул. Фурманова, дом 19, корпус 1 помещ. 2\n"
            "ООО «Банк Точка»\n"
            "БИК 044525104\n"
            "р/с 4070281050250031729\n"
            "к/с 30101810745374525104"
        )
        details_run = details_paragraph.add_run(details_text)
        set_font(details_run, 10)

        # Заголовок
        header = doc.add_paragraph(f"Исх. №{chat_id % 1000:03d} от {datetime.now().strftime('%d.%m.%Y')} г.")
        header.alignment = WD_ALIGN_PARAGRAPH.LEFT
        header.paragraph_format.space_after = Pt(10)
        header_run = header.runs[0]
        set_font(header_run, 12)

        # Заголовок КП
        title = doc.add_paragraph("Коммерческое предложение")
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER  # Центрирование надписи
        title.paragraph_format.space_after = Pt(10)
        title_run = title.runs[0]
        set_font(title_run, 14, bold=True)

        # Таблица
        materials_db = load_materials_db()
        templates_db = {item["name"]: item for item in materials_db.get("templates", []) if item["category"] == "Изделия"}

        # Подсчитываем количество строк (данные + заголовок + итог)
        total_rows = 1  # Заголовок
        for sheet in user_data[chat_id]["sheets"]:
            for item in user_data[chat_id]["products"].get(sheet, []):
                if item.get("name", "") in templates_db:
                    total_rows += 1
        total_rows += 1  # Итоговая строка

        table = doc.add_table(rows=total_rows, cols=5)
        table.autofit = False
        # Растягиваем таблицу на всю ширину страницы (180 мм после учета полей)
        table.columns[0].width = Mm(90)  # Наименование
        table.columns[1].width = Mm(25)  # Ед. изм.
        table.columns[2].width = Mm(25)  # Стоимость за ед.
        table.columns[3].width = Mm(15)  # Кол-во
        table.columns[4].width = Mm(25)  # Сумма

        # Заголовки таблицы
        headers = ["Наименование", "Ед. изм.", "Стоимость за ед.,\nруб.", "Кол-во", "Сумма, руб."]
        header_cells = table.rows[0].cells
        for i, header in enumerate(headers):
            p = header_cells[i].paragraphs[0]
            run = p.add_run(header)
            set_font(run, 12, bold=True)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if i != 0 else WD_ALIGN_PARAGRAPH.LEFT
            header_cells[i].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            try:
                set_cell_margins(header_cells[i], left=100, right=100, top=50, bottom=50)
                set_cell_borders(header_cells[i])
                header_cells[i]._element.tcPr.shd.set(qn('w:fill'), 'D3D3D3')  # Серый фон
            except Exception as e:
                logger.warning(f"Не удалось настроить стили для ячейки заголовка: {str(e)}")

        # Данные таблицы
        total_cost = 0
        row_idx = 1
        for sheet in user_data[chat_id]["sheets"]:
            for item in user_data[chat_id]["products"].get(sheet, []):
                name = item.get("name", "")
                if name not in templates_db:
                    logger.warning(f"Изделие '{name}' не найдено в materials.json")
                    continue
                unit = templates_db[name].get("unit", "шт")
                qty = item.get("quantity", 1)
                price = item.get("price_per_unit", 0)
                total = item.get("total_cost", qty * price)
                total_cost += total

                cells = table.rows[row_idx].cells
                cells[0].text = name
                cells[1].text = unit
                cells[2].text = f"{price:,.0f}".replace(',', '.')
                cells[3].text = str(qty)
                cells[4].text = f"{total:,.0f}".replace(',','.')
                

                for i in range(5):
                    p = cells[i].paragraphs[0]
                    run = p.runs[0]
                    set_font(run, 12)
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER if i != 0 else WD_ALIGN_PARAGRAPH.LEFT
                    cells[i].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                    try:
                        set_cell_margins(cells[i], left=100, right=100, top=50, bottom=50)
                        set_cell_borders(cells[i])
                    except Exception as e:
                        logger.warning(f"Не удалось настроить стили для ячейки данных: {str(e)}")
                row_idx += 1

        # Итоговая строка
        total_cells = table.rows[row_idx].cells
        total_cells[3].text = "Общая стоимость, руб."
        total_cells[4].text = f"{total_cost:,.0f}".replace(',','.')
        for i in [3, 4]:
            p = total_cells[i].paragraphs[0]
            run = p.runs[0]
            set_font(run, 12)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            total_cells[i].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            try:
                set_cell_margins(total_cells[i], left=100, right=100, top=50, bottom=50)
                set_cell_borders(total_cells[i])
            except Exception as e:
                logger.warning(f"Не удалось настроить стили для итоговой ячейки: {str(e)}")

        # Подпись
        signature = doc.add_paragraph("Генеральный директор ООО «СУП» Зонов Д.А.")
        signature.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        signature_run = signature.runs[0]
        set_font(signature_run, 12)

        doc.save(filename)
        logger.info(f"Коммерческое предложение DOCX сохранено: {filename}")
        return filename

    except Exception as e:
        logger.error(f"Ошибка при создании DOCX КП: {str(e)}", exc_info=True)
        raise